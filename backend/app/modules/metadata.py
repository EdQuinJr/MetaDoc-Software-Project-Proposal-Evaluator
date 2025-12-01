"""
Module 2: Metadata Extraction & Content Analysis

Implements SRS requirements:
- M2.UC01: Perform Metadata Extraction & Validation
- M2.UC02: Store Analysis Snapshot & Generate Report

Handles:
1. Extracting essential metadata fields from Google Docs and DOCX files
2. Parsing and processing full document text
3. Computing content statistics (word count, sentences, pages)
4. Validating document completeness
5. Generating analysis snapshots for version comparison
6. Preparing structured output for downstream modules
"""

import os
import json
import time
from datetime import datetime
from flask import Blueprint, request, jsonify, current_app
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter
import re

from app import db
from app.models import Submission, AnalysisResult, DocumentSnapshot, SubmissionStatus
from app.services.audit_service import AuditService

metadata_bp = Blueprint('metadata', __name__)

class MetadataExtractionService:
    """Service for extracting metadata and analyzing document content"""
    
    def __init__(self):
        pass
    
    @property
    def min_word_count(self):
        return current_app.config.get('MIN_DOCUMENT_WORDS', 50)
    
    @property
    def max_word_count(self):
        return current_app.config.get('MAX_DOCUMENT_WORDS', 15000)
    
    def extract_docx_metadata(self, file_path):
        """Extract metadata from DOCX file using python-docx and direct XML parsing"""
        metadata = {
            'author': 'Unavailable',
            'creation_date': None,
            'last_modified_date': None,
            'last_editor': 'Unavailable',
            'file_size': 0,
            'word_count': 0,
            'revision_count': 0,
            'application': 'Unknown'
        }
        
        try:
            # Get file size
            metadata['file_size'] = os.path.getsize(file_path)
            
            # Load document with python-docx
            doc = Document(file_path)
            
            # Extract basic properties
            core_props = doc.core_properties
            
            if core_props.author:
                metadata['author'] = core_props.author
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            
            if core_props.modified:
                metadata['last_modified_date'] = core_props.modified.isoformat()
            
            if core_props.last_modified_by:
                metadata['last_editor'] = core_props.last_modified_by
            
            if hasattr(core_props, 'revision') and core_props.revision:
                try:
                    metadata['revision_count'] = int(core_props.revision)
                except (ValueError, TypeError):
                    metadata['revision_count'] = 0
            
            # Try to extract additional metadata from XML
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Read app properties for more detailed metadata
                    if 'docProps/app.xml' in zip_file.namelist():
                        app_xml = zip_file.read('docProps/app.xml')
                        app_root = ET.fromstring(app_xml)
                        
                        # Find application name
                        for elem in app_root.iter():
                            if elem.tag.endswith('Application'):
                                metadata['application'] = elem.text or 'Unknown'
                                break
                        
                        # Find word count
                        for elem in app_root.iter():
                            if elem.tag.endswith('Words'):
                                try:
                                    metadata['word_count'] = int(elem.text or 0)
                                except (ValueError, TypeError):
                                    metadata['word_count'] = 0
                                break
                    
                    # Read core properties for additional timestamps
                    if 'docProps/core.xml' in zip_file.namelist():
                        core_xml = zip_file.read('docProps/core.xml')
                        core_root = ET.fromstring(core_xml)
                        
                        for elem in core_root.iter():
                            if elem.tag.endswith('created') and not metadata['creation_date']:
                                metadata['creation_date'] = elem.text
                            elif elem.tag.endswith('modified') and not metadata['last_modified_date']:
                                metadata['last_modified_date'] = elem.text
            
            except Exception as e:
                current_app.logger.warning(f"Could not extract extended metadata: {e}")
            
            return metadata, None
            
        except PackageNotFoundError:
            return None, "Invalid DOCX file format"
        except Exception as e:
            current_app.logger.error(f"Metadata extraction failed: {e}")
            return None, f"Metadata extraction error: {e}"
    
    def extract_document_text(self, file_path):
        """Extract full text content from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
            
            # Combine all text
            full_text = '\n'.join(paragraphs)
            if table_text:
                full_text += '\n' + '\n'.join(table_text)
            
            return full_text, None
            
        except Exception as e:
            current_app.logger.error(f"Text extraction failed: {e}")
            return None, f"Text extraction error: {e}"
    
    def compute_content_statistics(self, text):
        """Compute comprehensive content statistics"""
        if not text:
            return {
                'word_count': 0,
                'character_count': 0,
                'character_count_no_spaces': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'estimated_pages': 0,
                'average_words_per_sentence': 0,
                'average_sentence_length': 0
            }
        
        # Basic counts
        character_count = len(text)
        character_count_no_spaces = len(text.replace(' ', ''))
        
        # Word count (split by whitespace and filter empty strings)
        words = [word.strip() for word in text.split() if word.strip()]
        word_count = len(words)
        
        # Sentence count (split by sentence endings)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        sentence_count = len(sentences)
        
        # Paragraph count (split by double newlines or single newlines)
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        paragraph_count = len(paragraphs)
        
        # Estimated pages (assuming ~250 words per page)
        estimated_pages = max(1, round(word_count / 250))
        
        # Averages
        average_words_per_sentence = word_count / max(sentence_count, 1)
        average_sentence_length = character_count / max(sentence_count, 1)
        
        return {
            'word_count': word_count,
            'character_count': character_count,
            'character_count_no_spaces': character_count_no_spaces,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'estimated_pages': estimated_pages,
            'average_words_per_sentence': round(average_words_per_sentence, 2),
            'average_sentence_length': round(average_sentence_length, 2)
        }
    
    def validate_document_completeness(self, content_stats, text):
        """Validate document completeness according to SRS requirements"""
        warnings = []
        is_complete = True
        
        # Check minimum word count
        if content_stats['word_count'] < self.min_word_count:
            warnings.append(f"Document appears incomplete: only {content_stats['word_count']} words (minimum: {self.min_word_count})")
            is_complete = False
        
        # Check maximum word count
        if content_stats['word_count'] > self.max_word_count:
            warnings.append(f"Document exceeds maximum length: {content_stats['word_count']} words (maximum: {self.max_word_count})")
        
        # Check for empty or very short sentences
        if content_stats['sentence_count'] < 3:
            warnings.append("Document has very few sentences")
            is_complete = False
        
        # Check for placeholder text or templates
        placeholder_patterns = [
            r'\[.*\]',  # [Insert text here]
            r'lorem ipsum',
            r'placeholder',
            r'insert.*here',
            r'add.*content',
            r'write.*here'
        ]
        
        text_lower = text.lower() if text else ""
        for pattern in placeholder_patterns:
            if re.search(pattern, text_lower):
                warnings.append("Document may contain placeholder text")
                break
        
        # Check for repeated content (potential copy-paste issues)
        if text and len(text) > 100:
            # Simple check for repeated paragraphs
            paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 20]
            if len(paragraphs) != len(set(paragraphs)) and len(paragraphs) > 1:
                warnings.append("Document may contain repeated content")
        
        return is_complete, warnings
    
    def create_analysis_snapshot(self, submission, metadata, content_stats, text):
        """Create analysis snapshot for version comparison"""
        try:
            # Check for previous snapshots of the same document
            file_id = f"{submission.original_filename}_{submission.file_hash[:8]}"
            
            previous_snapshot = DocumentSnapshot.query.filter_by(
                file_id=file_id
            ).order_by(DocumentSnapshot.created_at.desc()).first()
            
            # Create new snapshot
            snapshot = DocumentSnapshot(
                file_id=file_id,
                submission_id=submission.id,
                word_count=content_stats['word_count'],
                file_hash=submission.file_hash,
                snapshot_timestamp=datetime.utcnow()
            )
            
            # Compare with previous snapshot if exists
            if previous_snapshot:
                word_count_change = content_stats['word_count'] - previous_snapshot.word_count
                if previous_snapshot.word_count > 0:
                    change_percentage = (word_count_change / previous_snapshot.word_count) * 100
                    snapshot.change_percentage = round(change_percentage, 2)
                    
                    # Flag major changes (≥50% as per SRS)
                    if abs(change_percentage) >= 50:
                        snapshot.major_changes = True
            
            db.session.add(snapshot)
            db.session.commit()
            
            return snapshot, None
            
        except Exception as e:
            current_app.logger.error(f"Failed to create snapshot: {e}")
            return None, f"Snapshot creation error: {e}"
    
    def generate_preliminary_report(self, submission, metadata, content_stats, text, is_complete, warnings):
        """Generate preliminary human-readable summary"""
        report = {
            'document_info': {
                'filename': submission.original_filename,
                'file_size_mb': round(submission.file_size / (1024 * 1024), 2),
                'submission_type': submission.submission_type,
                'submitted_at': submission.created_at.isoformat()
            },
            'metadata_summary': {
                'author': metadata.get('author', 'Unavailable'),
                'created': metadata.get('creation_date'),
                'last_modified': metadata.get('last_modified_date'),
                'last_editor': metadata.get('last_editor', 'Unavailable'),
                'revision_count': metadata.get('revision_count', 0)
            },
            'content_summary': {
                'word_count': content_stats['word_count'],
                'pages': content_stats['estimated_pages'],
                'sentences': content_stats['sentence_count'],
                'paragraphs': content_stats['paragraph_count'],
                'avg_words_per_sentence': content_stats['average_words_per_sentence']
            },
            'validation': {
                'is_complete': is_complete,
                'warnings': warnings,
                'meets_requirements': is_complete and content_stats['word_count'] >= self.min_word_count
            },
            'analysis_timestamp': datetime.utcnow().isoformat()
        }
        
        return report

# Initialize service
metadata_service = MetadataExtractionService()

@metadata_bp.route('/analyze/<submission_id>', methods=['POST'])
def analyze_submission(submission_id):
    """
    Perform complete metadata extraction and content analysis
    
    SRS Reference: M2.UC01 - Perform Metadata Extraction & Validation
    """
    start_time = time.time()
    
    try:
        # Get submission
        submission = Submission.query.filter_by(id=submission_id).first()
        
        if not submission:
            return jsonify({'error': 'Submission not found'}), 404
        
        if submission.status != SubmissionStatus.PENDING:
            return jsonify({'error': 'Submission already processed or processing'}), 400
        
        # Update status to processing
        submission.status = SubmissionStatus.PROCESSING
        submission.processing_started_at = datetime.utcnow()
        db.session.commit()
        
        # Log processing start
        AuditService.log_submission_event('processing_started', submission)
        
        # Extract metadata
        metadata, metadata_error = metadata_service.extract_docx_metadata(submission.file_path)
        
        if metadata_error:
            submission.status = SubmissionStatus.FAILED
            submission.error_message = metadata_error
            db.session.commit()
            return jsonify({'error': metadata_error}), 500
        
        # Extract text content
        text, text_error = metadata_service.extract_document_text(submission.file_path)
        
        if text_error:
            submission.status = SubmissionStatus.FAILED
            submission.error_message = text_error
            db.session.commit()
            return jsonify({'error': text_error}), 500
        
        # Compute content statistics
        content_stats = metadata_service.compute_content_statistics(text)
        
        # Validate document completeness
        is_complete, warnings = metadata_service.validate_document_completeness(content_stats, text)
        
        # Create analysis snapshot
        snapshot, snapshot_error = metadata_service.create_analysis_snapshot(
            submission, metadata, content_stats, text
        )
        
        if snapshot_error:
            current_app.logger.warning(f"Snapshot creation failed: {snapshot_error}")
        
        # Generate preliminary report
        preliminary_report = metadata_service.generate_preliminary_report(
            submission, metadata, content_stats, text, is_complete, warnings
        )
        
        # Calculate processing duration
        processing_duration = time.time() - start_time
        
        # Create or update analysis result
        analysis_result = AnalysisResult.query.filter_by(submission_id=submission.id).first()
        if not analysis_result:
            analysis_result = AnalysisResult(submission_id=submission.id)
        
        # Store results
        analysis_result.document_metadata = metadata
        analysis_result.content_statistics = content_stats
        analysis_result.document_text = text
        analysis_result.is_complete_document = is_complete
        analysis_result.validation_warnings = warnings
        analysis_result.processing_duration_seconds = round(processing_duration, 2)
        
        # Save to database
        db.session.add(analysis_result)
        
        # Update submission status
        submission.status = SubmissionStatus.WARNING if warnings else SubmissionStatus.COMPLETED
        submission.processing_completed_at = datetime.utcnow()
        
        db.session.commit()
        
        # Log completion
        AuditService.log_submission_event(
            'metadata_analysis_completed',
            submission,
            additional_metadata={
                'word_count': content_stats['word_count'],
                'is_complete': is_complete,
                'warning_count': len(warnings),
                'processing_duration': processing_duration
            }
        )
        
        return jsonify({
            'message': 'Metadata analysis completed successfully',
            'job_id': submission.job_id,
            'status': submission.status.value,
            'analysis_id': analysis_result.id,
            'preliminary_report': preliminary_report,
            'processing_info': {
                'duration_seconds': processing_duration,
                'started_at': submission.processing_started_at.isoformat(),
                'completed_at': submission.processing_completed_at.isoformat()
            }
        }), 200
        
    except Exception as e:
        # Update submission status on error
        if 'submission' in locals():
            submission.status = SubmissionStatus.FAILED
            submission.error_message = f"Processing error: {str(e)}"
            db.session.commit()
        
        current_app.logger.error(f"Metadata analysis error: {e}")
        return jsonify({'error': 'Internal processing error'}), 500

@metadata_bp.route('/result/<submission_id>', methods=['GET'])
def get_analysis_result(submission_id):
    """
    Get analysis results for a submission
    
    SRS Reference: M2.UC02 - Store Analysis Snapshot & Generate Report
    """
    try:
        submission = Submission.query.filter_by(id=submission_id).first()
        
        if not submission:
            return jsonify({'error': 'Submission not found'}), 404
        
        analysis_result = AnalysisResult.query.filter_by(submission_id=submission.id).first()
        
        if not analysis_result:
            return jsonify({
                'message': 'Analysis not yet completed',
                'status': submission.status.value,
                'job_id': submission.job_id
            }), 202
        
        # Get document snapshots for version comparison
        snapshots = DocumentSnapshot.query.filter_by(
            submission_id=submission.id
        ).order_by(DocumentSnapshot.created_at.desc()).all()
        
        response_data = {
            'submission': submission.to_dict(),
            'analysis_result': analysis_result.to_dict(),
            'snapshots': [
                {
                    'id': s.id,
                    'word_count': s.word_count,
                    'timestamp': s.snapshot_timestamp.isoformat(),
                    'major_changes': s.major_changes,
                    'change_percentage': s.change_percentage
                } for s in snapshots
            ]
        }
        
        # Log data access
        AuditService.log_data_access('view', submission.id, None)
        
        return jsonify(response_data)
        
    except Exception as e:
        current_app.logger.error(f"Result retrieval error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@metadata_bp.route('/reprocess/<submission_id>', methods=['POST'])
def reprocess_submission(submission_id):
    """Reprocess a submission's metadata analysis"""
    try:
        submission = Submission.query.filter_by(id=submission_id).first()
        
        if not submission:
            return jsonify({'error': 'Submission not found'}), 404
        
        # Reset status and clear previous results
        submission.status = SubmissionStatus.PENDING
        submission.processing_started_at = None
        submission.processing_completed_at = None
        submission.error_message = None
        
        # Remove existing analysis result
        existing_result = AnalysisResult.query.filter_by(submission_id=submission.id).first()
        if existing_result:
            db.session.delete(existing_result)
        
        db.session.commit()
        
        # Log reprocessing request
        AuditService.log_submission_event('reprocessing_requested', submission)
        
        # Trigger analysis
        return analyze_submission(submission_id)
        
    except Exception as e:
        current_app.logger.error(f"Reprocess error: {e}")
        return jsonify({'error': 'Internal server error'}), 500