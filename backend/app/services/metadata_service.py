"""
Metadata Extraction Service - Handles document metadata and content analysis

Extracted from api/metadata.py to follow proper service layer architecture.
"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from flask import current_app
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.core.extensions import db
from app.models import DocumentSnapshot


class MetadataService:
    """Service for extracting metadata and analyzing document content"""
    
    def __init__(self):
        pass
    
    @property
    def min_word_count(self):
        return current_app.config.get('MIN_DOCUMENT_WORDS', 50)
    
    @property
    def max_word_count(self):
        return current_app.config.get('MAX_DOCUMENT_WORDS', 15000)
    
    def extract_docx_metadata(self, file_path):
        """Extract metadata from DOCX file using python-docx and direct XML parsing"""
        metadata = {
            'author': 'Unavailable',
            'creation_date': None,
            'last_modified_date': None,
            'last_editor': 'Unavailable',
            'file_size': 0,
            'word_count': 0,
            'revision_count': 0,
            'application': 'Unknown'
        }
        
        try:
            # Get file size
            metadata['file_size'] = os.path.getsize(file_path)
            
            # Load document with python-docx
            doc = Document(file_path)
            
            # Extract basic properties
            core_props = doc.core_properties
            
            if core_props.author:
                metadata['author'] = core_props.author
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            
            if core_props.modified:
                metadata['last_modified_date'] = core_props.modified.isoformat()
            
            if core_props.last_modified_by:
                metadata['last_editor'] = core_props.last_modified_by
            
            if hasattr(core_props, 'revision') and core_props.revision:
                try:
                    metadata['revision_count'] = int(core_props.revision)
                except (ValueError, TypeError):
                    metadata['revision_count'] = 0
            
            # Try to extract additional metadata from XML
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Read app properties for more detailed metadata
                    if 'docProps/app.xml' in zip_file.namelist():
                        app_xml = zip_file.read('docProps/app.xml')
                        app_root = ET.fromstring(app_xml)
                        
                        # Find application name
                        for elem in app_root.iter():
                            if elem.tag.endswith('Application'):
                                metadata['application'] = elem.text or 'Unknown'
                                break
                        
                        # Find word count
                        for elem in app_root.iter():
                            if elem.tag.endswith('Words'):
                                try:
                                    metadata['word_count'] = int(elem.text or 0)
                                except (ValueError, TypeError):
                                    metadata['word_count'] = 0
                                break
                    
                    # Read core properties for additional timestamps
                    if 'docProps/core.xml' in zip_file.namelist():
                        core_xml = zip_file.read('docProps/core.xml')
                        core_root = ET.fromstring(core_xml)
                        
                        for elem in core_root.iter():
                            if elem.tag.endswith('created') and not metadata['creation_date']:
                                metadata['creation_date'] = elem.text
                            elif elem.tag.endswith('modified') and not metadata['last_modified_date']:
                                metadata['last_modified_date'] = elem.text
                            elif elem.tag.endswith('lastModifiedBy') and metadata['last_editor'] == 'Unavailable':
                                metadata['last_editor'] = elem.text
                                
                        current_app.logger.info(f"Manual XML Extraction - Creator: {metadata.get('author')}, LastEditor: {metadata.get('last_editor')}")
            
            except Exception as e:
                current_app.logger.warning(f"Could not extract extended metadata: {e}")
            
            # Fallback: If last_editor is still Unavailable, use author if available
            if metadata['last_editor'] == 'Unavailable' and metadata['author'] != 'Unavailable':
                metadata['last_editor'] = metadata['author']
            
            # If last_editor is 'python-docx', set to Unavailable (if author fallback didn't help)
            if 'python-docx' in metadata['last_editor']:
                metadata['last_editor'] = 'Unavailable'
            
            # Initialize contributors list for standard uploads
            metadata['contributors'] = []
            if metadata.get('author') and metadata['author'] != 'Unavailable':
                metadata['contributors'].append({
                    'name': metadata['author'], 
                    'role': 'Owner and Writer',
                    'date': metadata.get('creation_date')
                })
            
            if metadata.get('last_editor') and metadata['last_editor'] != 'Unavailable':
                # Avoid duplicates
                if not metadata['contributors'] or metadata['contributors'][0]['name'] != metadata['last_editor']:
                    metadata['contributors'].append({
                        'name': metadata['last_editor'], 
                        'role': 'Last Editor',
                        'date': metadata.get('last_modified_date')
                    })
            
            
            return metadata, None
            
        except PackageNotFoundError:
            return None, "Invalid DOCX file format"
        except Exception as e:
            current_app.logger.error(f"Metadata extraction failed: {e}")
            return None, f"Metadata extraction error: {e}"
    
    def extract_document_text(self, file_path):
        """Extract full text content from DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text.strip())
            
            # Combine all text
            full_text = '\n'.join(paragraphs)
            if table_text:
                full_text += '\n' + '\n'.join(table_text)
            
            return full_text, None
            
        except Exception as e:
            current_app.logger.error(f"Text extraction failed: {e}")
            return None, f"Text extraction error: {e}"
    
    def compute_content_statistics(self, text):
        """Compute comprehensive content statistics"""
        if not text:
            return {
                'word_count': 0,
                'character_count': 0,
                'character_count_no_spaces': 0,
                'sentence_count': 0,
                'paragraph_count': 0,
                'estimated_pages': 0,
                'average_words_per_sentence': 0,
                'average_sentence_length': 0
            }
        
        # Basic counts
        character_count = len(text)
        character_count_no_spaces = len(text.replace(' ', ''))
        
        # Word count (split by whitespace and filter empty strings)
        words = [word.strip() for word in text.split() if word.strip()]
        word_count = len(words)
        
        # Sentence count (split by sentence endings)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        sentence_count = len(sentences)
        
        # Paragraph count (split by double newlines or single newlines)
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        paragraph_count = len(paragraphs)
        
        # Estimated pages (assuming ~250 words per page)
        estimated_pages = max(1, round(word_count / 250))
        
        # Averages
        average_words_per_sentence = word_count / max(sentence_count, 1)
        average_sentence_length = character_count / max(sentence_count, 1)
        
        return {
            'word_count': word_count,
            'character_count': character_count,
            'character_count_no_spaces': character_count_no_spaces,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'estimated_pages': estimated_pages,
            'page_count': estimated_pages,  # Alias for frontend compatibility
            'average_words_per_sentence': round(average_words_per_sentence, 2),
            'average_sentence_length': round(average_sentence_length, 2)
        }
    
    def validate_document_completeness(self, content_stats, text):
        """Validate document completeness according to SRS requirements"""
        # Validation warnings disabled by user request
        warnings = []
        is_complete = True
        
        # All validation checks removed to prevent warnings
        
        return is_complete, warnings
    
    def create_analysis_snapshot(self, submission, metadata, content_stats, text):
        """Create analysis snapshot for version comparison"""
        try:
            # Check for previous snapshots of the same document
            file_id = f"{submission.original_filename}_{submission.file_hash[:8]}"
            
            previous_snapshot = DocumentSnapshot.query.filter_by(
                file_id=file_id
            ).order_by(DocumentSnapshot.created_at.desc()).first()
            
            # Create new snapshot
            snapshot = DocumentSnapshot(
                file_id=file_id,
                submission_id=submission.id,
                word_count=content_stats['word_count'],
                file_hash=submission.file_hash,
                snapshot_timestamp=datetime.utcnow()
            )
            
            # Compare with previous snapshot if exists
            if previous_snapshot:
                word_count_change = content_stats['word_count'] - previous_snapshot.word_count
                if previous_snapshot.word_count > 0:
                    change_percentage = (word_count_change / previous_snapshot.word_count) * 100
                    snapshot.change_percentage = round(change_percentage, 2)
                    
                    # Flag major changes (≥50% as per SRS)
                    if abs(change_percentage) >= 50:
                        snapshot.major_changes = True
            
            db.session.add(snapshot)
            db.session.commit()
            
            return snapshot, None
            
        except Exception as e:
            current_app.logger.error(f"Failed to create snapshot: {e}")
            return None, f"Snapshot creation error: {e}"
    
    def generate_preliminary_report(self, submission, metadata, content_stats, text, is_complete, warnings):
        """Generate preliminary human-readable summary"""
        report = {
            'document_info': {
                'filename': submission.original_filename,
                'file_size_mb': round(submission.file_size / (1024 * 1024), 2),
                'submission_type': submission.submission_type,
                'submitted_at': submission.created_at.isoformat()
            },
            'metadata_summary': {
                'author': metadata.get('author', 'Unavailable'),
                'created': metadata.get('creation_date'),
                'last_modified': metadata.get('last_modified_date'),
                'last_editor': metadata.get('last_editor', 'Unavailable'),
                'revision_count': metadata.get('revision_count', 0)
            },
            'content_summary': {
                'word_count': content_stats['word_count'],
                'pages': content_stats['estimated_pages'],
                'sentences': content_stats['sentence_count'],
                'paragraphs': content_stats['paragraph_count'],
                'avg_words_per_sentence': content_stats['average_words_per_sentence']
            },
            'validation': {
                'is_complete': is_complete,
                'warnings': warnings,
                'meets_requirements': is_complete and content_stats['word_count'] >= self.min_word_count
            },
            'analysis_timestamp': datetime.utcnow().isoformat()
        }
        
        return report
